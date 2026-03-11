import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

# Set Default Font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

def add_header_center(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)

def add_text_center(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)

def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    add_horizontal_line(p)

def add_entry(title, loc_or_role, date, desc=None, bullets=[]):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(5.8)
    table.columns[1].width = Inches(1.5)
    
    row = table.rows[0]
    p_left = row.cells[0].paragraphs[0]
    p_left.paragraph_format.space_after = Pt(0)
    
    # BOLD the Title (Event Name / Project Name / Company)
    r1 = p_left.add_run(f"{title}")
    r1.bold = True
    
    if loc_or_role:
        p_left.add_run(f" - {loc_or_role}")
        
    p_right = row.cells[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(0)
    p_right.add_run(date)
    
    if desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_after = Pt(2)
        r_desc = p_desc.add_run(desc)
        r_desc.italic = True
        
    for b in bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(0)
        p_b.paragraph_format.left_indent = Inches(0.25)
        p_b.add_run(b)

# ─── HEADER ───
add_header_center("M. RIZKY PRATAMA")
add_text_center("+62 882 1777 6891 | mrizky2pratama3@gmail.com | Sidoarjo, Jawa Timur | r1zky.vercel.app")

p_summ = doc.add_paragraph("Mechatronics Engineering student in 6th semester at PENS with a robust technical background and extensive experience in event organization, technical projects, and leadership. Highly active in organizations with a proven track record handling varying technical and non-technical divisions for both local and international scale events.")
p_summ.paragraph_format.space_before = Pt(8)

# ─── EDUCATION ───
add_section_header("Education Level")
add_entry("Politeknik Elektronika Negeri Surabaya (PENS)", "Surabaya, Indonesia", "2022 - Present", 
          "Bachelor of Applied Science in Mechatronics Engineering (Semester 6)", 
          ["Active member of IEEE SB PENS and GDGoC PENS.", "Hosted and managed various technical, leadership, and community events."])
add_entry("SMK Migas Cepu", "Cepu, Indonesia", "2020 - 2023", 
          "Industrial Electronics", [])
add_entry("SMPN 2 Sidoarjo", "Sidoarjo, Indonesia", "2017 - 2020", 
          "Junior High School", [])

# ─── INTERNSHIPS ───
add_section_header("Internship Experience")
add_entry("PT PAL Indonesia", "Design Division", "Jan - May 2026", 
          None, ["Field practice and design division operations."])
add_entry("PT Semen Indonesia (Persero) Tbk.", "Maintenance Division", "Jul - Sep 2022", 
          None, ["Praktik Kerja Lapangan (Field Internship)."])

# ─── ORGANISATIONS & AFFILIATIONS ───
add_section_header("Organisations & Affiliations")
add_entry("IEEE Student Branch PENS", "Program & Activities", "2025 - 2026", None, [])
add_entry("GDGoC PENS", "Event Manager", "2025 - 2026", None, [])
add_entry("BluAmbassador East Java", "Member", "2025 - 2027", None, [])
add_entry("Komunitas Sahabat Bahasa PENS", "Daily Operations Committee Staff", "2025 - 2027", None, [])

# ─── FIELD & EVENT EXPERIENCE ───
add_section_header("Field & Event Experience")

# Project Leader
add_entry("Data Science 101: Taking Your First Steps in the World of Data", "PIC", "Mar 6, 2026", 
          None, ["Person-in-Charge of the data science 101 event, coordinating speakers and flow."])
add_entry("Career Roadmap: Navigating the AI Engineering Landscape", "PIC", "Feb 27, 2026", 
          None, ["Person-in-Charge of the career roadmap event, coordinating speakers and flow."])
add_entry("International Webinar 'Unlock Global Potential'", "Project Leader (Ketua Pelaksana)", "Oct 18, 2025", 
          None, ["Led an international-scale webinar event coordinating stakeholders & global speakers."])

# MC
add_entry("Final Project Competition 2025", "MC", "Jul 1-3, 2025", 
          None, ["Hosted the main stage for final project presentations and guided audience through 3-day competition."])
add_entry("Navigasi Karir", "MC", "Jun 14, 2025", 
          None, ["Facilitated a career navigation event & audience interaction."])
add_entry("LKMM TD PENS 2025", "MC", "May 15-17, 2025", 
          None, ["Hosted official student leadership training."])
add_entry("FPWP Mekatronika 2024", "MC", "Oct 4, 2024", 
          None, ["Hosted FPWP (Ekspresi Kreasi Teknik Mekatronika) event."])
add_entry("Serah Terima Jabatan Mekatronika 2024/2025", "MC", "Sep 19, 2024", 
          None, ["Hosted the official handover ceremony for Mekatronika student leadership 2024/2025."])
add_entry("PKKMB × TECHNOGEAR 2024", "MC", "Aug 12-17, 2024", 
          None, ["Facilitated official student orientation event."])

# Event Division (Stage, Sound, Lighting, Content)
add_entry("Reuni 40th SMPN 2 Surabaya", "Stage Director", "Dec 7, 2025", 
          None, ["Managed stage flow & show timing, coordinated performers & crew."])
add_entry("SSHM RUN 2025", "Sound Director", "Aug 1-3, 2025", 
          None, ["Managed audio quality & sound system for the run event."])
add_entry("Berbagi Kebahagiaan 2025 — FKMPI Jawa Timur", "Timekeeper", "Mar 25, 2025", 
          None, ["Ensured event timeline adhered precisely to the rundown."])
add_entry("Pre-Competition Concert — PENS Students Choir", "Lighting Director, Scenario & LPJ", "Aug 31, 2024", 
          None, ["Designed stage lighting, wrote event scenario and final accountability report (LPJ)."])
add_entry("Final Project Competition 2024", "Operator", "Jul 9-11, 2024", 
          None, ["Operated technical event systems."])

# Support Divisions
add_entry("Flutter Fusion Conference", "Volunteer", "Feb 21, 2026", 
          None, ["Managed equipment setup and logistics for the Flutter community conference."])
add_entry("Chibicon 9", "Safety Division Member", "Dec 27-28, 2025", 
          None, ["Crowd & visitor flow control, maintained event area order."])
add_entry("Surabaya Fun Run 2025", "Health Division Member", "Oct 5, 2025", 
          None, ["Provided first aid for run participants and managed medical post operations."])
add_entry("Chibicon! 7", "Safety Division Member", "Jun 21-22, 2025", 
          None, ["Crowd & visitor flow control, maintained event area order."])
add_entry("Chibicon!", "Equipment Staff", "Apr 20-21, 2024", 
          None, ["Managed equipment setup, logistics & teardown for the event."])
add_entry("LKMM Pra Tingkat Dasar", "Liaison Officer", "Oct 25-27, 2024", 
          None, ["Accompanied & assisted participants throughout the event."])

# ─── PROJECTS ───
add_section_header("Work & Projects")
add_entry("Balancing Bot", "Inventor, Eagle, Proteus", "Aug - Dec 2025", 
          None, ["Self-balancing robot utilizing PID control and sensor fusion for real-time dynamic balance."])
add_entry("CNC Plotter", "Visual Studio Code, Inventor", "Aug - Dec 2025", 
          None, ["Computer-controlled 2D plotting machine that draws patterns using G-code from digital designs."])
add_entry("PID Analog", "Proteus, Inventor", "May - Jun 2025", 
          None, ["Analog PID-based DC motor speed control system maintaining setpoint via proportional-integral-derivative control."])
add_entry("RTSS — Robot Tangan Satu Sendi", "Proteus, Inventor", "Feb - May 2025", 
          None, ["Single-joint robotic arm system simulated and designed for mechanical analysis and control."])
add_entry("CNC Machine Control System", "Autodesk Inventor", "Dec 2024", 
          None, ["Innovative CNC machine control system integrated with Autodesk Inventor."])
add_entry("Scoring Board", "Electronics, PCB Design", "Aug - Dec 2024", 
          None, ["Electronic scoring board with dual 7-segment displays and potentiometer-based score control."])
add_entry("Workshop 3D: Making Urban Toys", "Blender", "Sep - Oct 2024", 
          None, ["3D design & printing workshop focused on crafting urban-themed toys."])
add_entry("CNC-Based Project", "Winmax, CNC Haas Mill VF3", "Feb - May 2024", 
          None, ["Precision component design using CNC technology with WinMax & CAD/CAM."])

# ─── SKILLS & ACHIEVEMENTS ───
add_section_header("Skills & Achievements")
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Achievements: ").bold = True
p1.add_run("Finalist — Master of Ceremony Category (PENS ELITE Competition), Very Good — English Skill (Kampung Inggris Online), Best Social Technology Development Program (IEEE Indonesia Section).")

p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Hard Skills: ").bold = True
p2.add_run("PCB Design, Proteus, Autodesk Inventor, Blender, CNC Haas Mill VF3, Winmax, CAD/CAM, PID Control, Visual Studio Code.")

p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Soft Skills: ").bold = True
p3.add_run("Leadership, Time Management, Teamwork, Communication, Critical Thinking, Problem Solving, Creativity, Innovation.")

p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Languages: ").bold = True
p4.add_run("Bahasa Indonesia (Native), English (Proficient).")

doc.save('M_Rizky_Pratama_CV.docx')
print("Document updated successfully.")
